"""Parse DOCX and PDF files into raw text and structural data."""
import hashlib
from pathlib import Path
from docx import Document
import pdfplumber


def file_hash(path: str) -> str:
    h = hashlib.sha256(Path(path).read_bytes()).hexdigest()
    return f"sha256:{h[:16]}"


def parse_docx(path: str) -> dict:
    """Extract text, tables, styles, and colors from a DOCX file."""
    doc = Document(path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    tables = []
    for i, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        # Detect background colors on cells
        colors = set()
        for row in table.rows:
            for cell in row.cells:
                shd = cell._tc.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
                )
                if shd is not None:
                    fill = shd.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"
                    )
                    if fill and fill.upper() not in ("AUTO", "FFFFFF", ""):
                        colors.add(fill.upper())
        tables.append({"index": i, "rows": rows, "colors": list(colors)})

    styles = {}
    for s in doc.styles:
        try:
            if s.font and s.font.name:
                styles[s.name] = {"font": s.font.name, "size": str(s.font.size)}
        except Exception:
            pass

    # Collect all unique colors from document
    all_colors: set[str] = set()
    for t in tables:
        all_colors.update(t["colors"])
    for para in doc.paragraphs:
        for run in para.runs:
            try:
                if run.font.color.rgb:
                    all_colors.add(str(run.font.color.rgb).upper())
            except Exception:
                pass

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "styles": styles,
        "detected_colors": list(all_colors),
        "hash": file_hash(path),
    }


def parse_pdf(path: str) -> dict:
    """Extract text and tables from a text-based PDF."""
    pages = []
    tables = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages.append({"page": i + 1, "text": text.strip()})
            for tbl in page.extract_tables() or []:
                tables.append({"page": i + 1, "rows": tbl})

    full_text = "\n".join(p["text"] for p in pages)
    return {
        "pages": pages,
        "tables": tables,
        "full_text": full_text,
        "hash": file_hash(path),
    }


def ingest(path: str) -> dict:
    """Auto-detect format and parse the file."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        data = parse_pdf(path)
        data["format"] = "pdf"
    else:
        data = parse_docx(path)
        data["format"] = "docx"
    return data
