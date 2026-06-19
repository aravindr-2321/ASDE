"""Blueprint-driven DOCX assembler using python-docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from pathlib import Path
from ase.schemas.models import TemplateBlueprint, ContentModel, SubjectContent


def _hex_to_rgb(h: str) -> RGBColor:
    h = h.lstrip("#").upper()
    if len(h) != 6:
        return RGBColor(0x66, 0x00, 0x66)
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return RGBColor(r, g, b)


def _shade_cell(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper().lstrip("#"))
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold=False, color: str = "", font_size: int = 10,
                   align=WD_ALIGN_PARAGRAPH.LEFT, font: str = "Calibri"):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = _hex_to_rgb(color)


def _header_row(table, texts: list[str], bg: str, fg: str = "FFFFFF",
                font: str = "Calibri", size: int = 10):
    row = table.rows[0]
    for i, cell in enumerate(row.cells):
        txt = texts[i] if i < len(texts) else ""
        _shade_cell(cell, bg)
        _set_cell_text(cell, txt, bold=True, color=fg, font_size=size,
                       align=WD_ALIGN_PARAGRAPH.CENTER, font=font)


def _add_banner_row(table, label: str, bg: str, fg: str = "FFFFFF", merge=True):
    """Add a full-width banner row to a table."""
    row = table.add_row()
    if merge and len(row.cells) > 1:
        row.cells[0].merge(row.cells[-1])
    cell = row.cells[0]
    _shade_cell(cell, bg)
    _set_cell_text(cell, label, bold=True, color=fg, font_size=10,
                   align=WD_ALIGN_PARAGRAPH.LEFT)


def build_docx(
    blueprint: TemplateBlueprint,
    content: ContentModel,
    output_path: str,
) -> str:
    """Assemble a complete university syllabus DOCX from blueprint + content."""
    doc = Document()

    ct = blueprint.color_tokens
    primary = ct.get("table_header", "660066")
    accent = ct.get("title_box", "F7931D")
    heading_color = ct.get("heading", "6F1952")
    label_dict = blueprint.label_dictionary
    obj_code = label_dict.get("objective_code", "CLO")
    default_font = blueprint.typography.get("default_font", "Calibri")
    body_font = blueprint.typography.get("body_font", "Times New Roman")

    # ── Page setup ───────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Title page ───────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{blueprint.university_id.upper()}\nSyllabus Document")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = default_font
    run.font.color.rgb = _hex_to_rgb(heading_color)
    doc.add_paragraph(f"Program: {content.program}")
    doc.add_paragraph(f"Semester: {content.semester}")
    doc.add_page_break()

    # ── Credit Structure Table ────────────────────────────────────────────────
    p = doc.add_paragraph("CREDIT STRUCTURE")
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = _hex_to_rgb(heading_color)

    col_labels = ["S.No", "Subject Name", "Code", "L", "T", "P", "Credits", "Marks"]
    tbl = doc.add_table(rows=1, cols=len(col_labels))
    tbl.style = "Table Grid"
    _header_row(tbl, col_labels, primary, "FFFFFF", default_font)

    for i, subj in enumerate(content.subjects, 1):
        row = tbl.add_row()
        vals = [
            str(i), subj.name, subj.code or "—",
            *(subj.ltp.split("-") if subj.ltp and "-" in subj.ltp else ["—", "—", "—"]),
            subj.credits or "—", subj.marks or "—",
        ]
        for j, cell in enumerate(row.cells):
            _set_cell_text(cell, vals[j] if j < len(vals) else "—",
                           font_size=9, font=body_font)

    doc.add_page_break()

    # ── Subject Detail Pages ──────────────────────────────────────────────────
    for subj in content.subjects:
        _add_subject_page(doc, subj, blueprint, content.program,
                          content.semester, primary, accent, heading_color,
                          default_font, body_font, obj_code, label_dict)
        doc.add_page_break()

    # ── References Section ────────────────────────────────────────────────────
    if blueprint.has_sections.get("textbooks") or blueprint.has_sections.get("references"):
        _add_references_section(doc, content.subjects, primary, heading_color, default_font, body_font)

    doc.save(output_path)
    return output_path


def _add_subject_page(doc, subj: SubjectContent, bp: TemplateBlueprint,
                      program, semester, primary, accent, heading_color,
                      default_font, body_font, obj_code, label_dict):
    """Add a table-packed detail page for one subject."""
    tbl = doc.add_table(rows=0, cols=1)
    tbl.style = "Table Grid"

    # Course scheme header
    _add_banner_row(tbl, f"COURSE: {subj.name}  |  {program}  |  Semester {semester}", primary, "FFFFFF")

    # Level row
    level_label = label_dict.get("level_row_label", "Level (UG/PG) and NcRF")
    row = tbl.add_row()
    _shade_cell(row.cells[0], "F2F2F2")
    _set_cell_text(row.cells[0], f"{level_label}: Undergraduate (UG)", font_size=9, font=body_font)

    # CLO banner + content
    obj_banner = label_dict.get("objectives_banner", f"Course Learning Objectives ({obj_code}):")
    _add_banner_row(tbl, obj_banner, accent, "FFFFFF")

    objs = subj.generated_objectives or subj.objectives
    obj_text = "\n".join(objs) if objs else "To be determined."
    row = tbl.add_row()
    _set_cell_text(row.cells[0], obj_text, font_size=9, font=body_font)

    # CO banner + content
    co_banner = label_dict.get("outcomes_banner", "Course Outcomes (CO):")
    _add_banner_row(tbl, co_banner, accent, "FFFFFF")

    outs = subj.generated_outcomes or subj.outcomes
    co_text = "\n".join(outs) if outs else "To be determined."
    row = tbl.add_row()
    _set_cell_text(row.cells[0], co_text, font_size=9, font=body_font)

    # Course content banner
    content_banner = label_dict.get("content_banner", "Course Content:")
    _add_banner_row(tbl, content_banner, primary, "FFFFFF")

    for mod in subj.modules:
        # Module header sub-row
        row = tbl.add_row()
        _shade_cell(row.cells[0], "E8E8E8")
        _set_cell_text(row.cells[0], f"{mod.label}: {mod.title}", bold=True, font_size=9, font=default_font)
        # Topics row
        row = tbl.add_row()
        _set_cell_text(row.cells[0], mod.topics, font_size=9, font=body_font)

    # CO-PO mapping
    copo_banner = label_dict.get("copo_banner", "CO-PO Mapping:")
    _add_banner_row(tbl, copo_banner, primary, "FFFFFF")
    _add_copo_table(doc, tbl, subj, primary, body_font)


def _add_copo_table(doc, parent_tbl, subj: SubjectContent, primary, font):
    """Append CO-PO matrix into a new row spanning the parent table."""
    copo = subj.copo
    if not copo:
        row = parent_tbl.add_row()
        _set_cell_text(row.cells[0], "CO-PO mapping not available.", font_size=9, font=font)
        return

    po_count = copo.get("po_count", 12)
    scale = copo.get("scale", ["H", "M", "L", "-"])
    rows_data = copo.get("rows", [])

    # Build as text in the cell (table-in-table is complex; use simple layout)
    po_headers = "\t".join(f"PO{i+1}" for i in range(po_count))
    lines = [f"CO\t{po_headers}"]
    for i, row_vals in enumerate(rows_data):
        line = f"CO{i+1}\t" + "\t".join(str(v) for v in row_vals)
        lines.append(line)
    lines.append(f"Scale: {' / '.join(scale)}")

    row = parent_tbl.add_row()
    _set_cell_text(row.cells[0], "\n".join(lines), font_size=8, font=font)


def _add_references_section(doc, subjects: list[SubjectContent],
                             primary, heading_color, default_font, body_font):
    p = doc.add_paragraph("TEXTBOOKS AND REFERENCES")
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = _hex_to_rgb(heading_color)

    for subj in subjects:
        doc.add_paragraph(subj.name, style="Heading 3")
        tb = subj.generated_references or subj.textbooks
        if tb:
            doc.add_paragraph("Textbooks:", style="Normal").runs[0].bold = True
            for t in tb:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(t).font.size = Pt(9)

        refs = subj.references
        if refs:
            doc.add_paragraph("References:", style="Normal").runs[0].bold = True
            for r in refs:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(r).font.size = Pt(9)
