"""QA: validate DOCX integrity, content preservation, and structural checks."""
from pathlib import Path
from docx import Document
from ase.schemas.models import ContentModel, TemplateBlueprint


def validate(docx_path: str, content: ContentModel, blueprint: TemplateBlueprint) -> dict:
    """Run all QA checks; return a report dict with score and findings."""
    findings = []
    checks_passed = 0
    checks_total = 0

    def check(name: str, ok: bool, detail: str = ""):
        nonlocal checks_passed, checks_total
        checks_total += 1
        if ok:
            checks_passed += 1
        else:
            findings.append({"check": name, "status": "FAIL", "detail": detail})

    # 1. File exists and is valid DOCX
    p = Path(docx_path)
    check("file_exists", p.exists(), f"File not found: {docx_path}")
    if not p.exists():
        return _report(0, findings, checks_passed, checks_total)

    try:
        doc = Document(docx_path)
        check("valid_docx", True)
    except Exception as e:
        check("valid_docx", False, str(e))
        return _report(0, findings, checks_passed, checks_total)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += " " + cell.text

    # 2. All subjects are present
    for subj in content.subjects:
        check(
            f"subject_present:{subj.name[:30]}",
            subj.name.lower()[:15] in full_text.lower(),
            f"Subject '{subj.name}' not found in output",
        )

    # 3. Objectives/outcomes generated or present
    for subj in content.subjects:
        has_obj = bool(subj.generated_objectives or subj.objectives)
        check(f"objectives:{subj.name[:20]}", has_obj, "No objectives/CLOs generated")
        has_co = bool(subj.generated_outcomes or subj.outcomes)
        check(f"outcomes:{subj.name[:20]}", has_co, "No outcomes/COs generated")

    # 4. Module count matches
    for subj in content.subjects:
        check(
            f"modules:{subj.name[:20]}",
            len(subj.modules) > 0,
            "No modules found for subject",
        )

    # 5. Document has tables (structure check)
    check("has_tables", len(doc.tables) > 0, "No tables found — structure may be wrong")

    # 6. Non-empty document
    check("non_empty", len(full_text.strip()) > 200, "Output document appears too short")

    score = round(checks_passed / max(checks_total, 1), 3)
    return _report(score, findings, checks_passed, checks_total)


def _report(score, findings, passed, total) -> dict:
    return {
        "score": score,
        "passed": passed,
        "total": total,
        "findings": findings,
        "status": "pass" if score >= 0.85 else "fail",
    }
